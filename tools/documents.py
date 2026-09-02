import sys
import json
import base64
import io
import requests

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image, PageBreak
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

COR_PRIMARIA = colors.HexColor("#1A2B4C")
COR_SECUNDARIA = colors.HexColor("#3E6B9E")
COR_DESTAQUE = colors.HexColor("#E8A33D")
COR_TEXTO = colors.HexColor("#2B2B2B")
COR_CINZA_CLARO = colors.HexColor("#F2F4F7")


def get_styles():
    styles = getSampleStyleSheet()
    return {
        "titulo": ParagraphStyle("Titulo", parent=styles["Title"], fontName="Helvetica-Bold",
                                  fontSize=24, leading=28, textColor=COR_PRIMARIA, spaceAfter=10),
        "subtitulo": ParagraphStyle("Subtitulo", parent=styles["Normal"], fontName="Helvetica",
                                     fontSize=13, leading=17, textColor=COR_SECUNDARIA, spaceAfter=18),
        "heading": ParagraphStyle("Heading", parent=styles["Heading1"], fontName="Helvetica-Bold",
                                   fontSize=15, leading=19, textColor=COR_PRIMARIA,
                                   spaceBefore=16, spaceAfter=8),
        "corpo": ParagraphStyle("Corpo", parent=styles["Normal"], fontName="Helvetica",
                                 fontSize=10.5, leading=15.5, textColor=COR_TEXTO, spaceAfter=8),
        "bullet": ParagraphStyle("Bullet", parent=styles["Normal"], fontName="Helvetica",
                                  fontSize=10.5, leading=15, leftIndent=14, spaceAfter=5,
                                  textColor=COR_TEXTO),
    }


def linha_divisoria():
    return HRFlowable(width="100%", thickness=1.5, color=COR_DESTAQUE, spaceBefore=4, spaceAfter=14)


def rodape_pagina(canvas_obj, doc):
    canvas_obj.saveState()
    canvas_obj.setStrokeColor(COR_CINZA_CLARO)
    canvas_obj.line(2 * cm, 1.7 * cm, A4[0] - 2 * cm, 1.7 * cm)
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(colors.grey)
    canvas_obj.drawRightString(A4[0] - 2 * cm, 1.2 * cm, f"Página {doc.page}")
    canvas_obj.restoreState()


def baixar_imagem(url, max_width_cm=15):
    try:
        resp = requests.get(url, timeout=8)
        resp.raise_for_status()
        img = Image(io.BytesIO(resp.content))
        largura_max = max_width_cm * cm
        proporcao = img.imageHeight / img.imageWidth
        img.drawWidth = largura_max
        img.drawHeight = largura_max * proporcao
        return img
    except Exception:
        return None


def montar_tabela(headers, rows):
    dados = [headers] + rows
    largura_col = (A4[0] - 4 * cm) / len(headers)
    tabela = Table(dados, colWidths=[largura_col] * len(headers))
    tabela.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), COR_PRIMARIA),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, COR_CINZA_CLARO]),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#DDDDDD")),
    ]))
    return tabela


def create_pdf_structured(params):
    """
    params = {
      "title": str,
      "subtitle": str (opcional),
      "sections": [
        {
          "heading": str,
          "paragraphs": [str, ...],
          "bullet_list": [str, ...],
          "image_url": str (opcional),
          "table": {"headers": [...], "rows": [[...], ...]}
        }
      ]
    }
    Retorna: dict com "pdf_base64"
    """
    s = get_styles()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                             topMargin=2.2 * cm, bottomMargin=2.2 * cm,
                             leftMargin=2 * cm, rightMargin=2 * cm,
                             title=params.get("title", "Documento"))

    story = [Paragraph(params["title"], s["titulo"])]
    if params.get("subtitle"):
        story.append(Paragraph(params["subtitle"], s["subtitulo"]))
    story.append(linha_divisoria())

    for sec in params.get("sections", []):
        if sec.get("heading"):
            story.append(Paragraph(sec["heading"], s["heading"]))

        for p in sec.get("paragraphs", []):
            story.append(Paragraph(p, s["corpo"]))

        for item in sec.get("bullet_list", []):
            story.append(Paragraph(f"• {item}", s["bullet"]))

        if sec.get("image_url"):
            img = baixar_imagem(sec["image_url"])
            if img:
                story.append(Spacer(1, 6))
                story.append(img)
                story.append(Spacer(1, 6))

        if sec.get("table"):
            t = sec["table"]
            story.append(Spacer(1, 6))
            story.append(montar_tabela(t["headers"], t["rows"]))
            story.append(Spacer(1, 10))

    doc.build(story, onFirstPage=rodape_pagina, onLaterPages=rodape_pagina)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return {"pdf_base64": base64.b64encode(pdf_bytes).decode("utf-8")}


def create_pdf(params):
    """
    Versão simplificada: título + lista de parágrafos/bullets direto,
    sem precisar estruturar em "sections". Internamente reaproveita
    create_pdf_structured com uma única seção.
    """
    secao = {
        "paragraphs": params.get("paragraphs", []),
        "bullet_list": params.get("bullet_list", []),
        "image_url": params.get("image_url"),
        "table": params.get("table"),
    }
    return create_pdf_structured({
        "title": params["title"],
        "subtitle": params.get("subtitle"),
        "sections": [secao],
    })


def create_docx(params):
    """
    params = {
      "title": str,
      "subtitle": str (opcional),
      "sections": [
        {
          "heading": str,
          "paragraphs": [str, ...],
          "bullet_list": [str, ...],
          "image_url": str (opcional),
          "table": {"headers": [...], "rows": [[...], ...]}
        }
      ]
    }
    Retorna: dict com "docx_base64"
    """
    document = Document()

    titulo = document.add_heading(params["title"], level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if params.get("subtitle"):
        sub = document.add_paragraph(params["subtitle"])
        sub.runs[0].font.size = Pt(13)
        sub.runs[0].font.color.rgb = RGBColor(0x3E, 0x6B, 0x9E)

    for sec in params.get("sections", []):
        if sec.get("heading"):
            document.add_heading(sec["heading"], level=1)

        for p in sec.get("paragraphs", []):
            document.add_paragraph(p)

        for item in sec.get("bullet_list", []):
            document.add_paragraph(item, style="List Bullet")

        if sec.get("image_url"):
            try:
                resp = requests.get(sec["image_url"], timeout=8)
                resp.raise_for_status()
                document.add_picture(io.BytesIO(resp.content), width=Cm(14))
            except Exception:
                pass

        if sec.get("table"):
            t = sec["table"]
            tabela = document.add_table(rows=1, cols=len(t["headers"]))
            tabela.style = "Light Grid Accent 1"
            hdr_cells = tabela.rows[0].cells
            for i, h in enumerate(t["headers"]):
                hdr_cells[i].text = h
            for row in t["rows"]:
                cells = tabela.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)

    buffer = io.BytesIO()
    document.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return {"docx_base64": base64.b64encode(docx_bytes).decode("utf-8")}


def merge_pdfs(params):
    from pypdf import PdfReader, PdfWriter
    writer = PdfWriter()
    for pdf_b64 in params["pdfs_base64"]:
        reader = PdfReader(io.BytesIO(base64.b64decode(pdf_b64)))
        for page in reader.pages:
            writer.add_page(page)
    buffer = io.BytesIO()
    writer.write(buffer)
    return {"pdf_base64": base64.b64encode(buffer.getvalue()).decode("utf-8")}


def split_pdf_pages(params):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(io.BytesIO(base64.b64decode(params["pdf_base64"])))
    writer = PdfWriter()
    for n in params["page_numbers"]:
        writer.add_page(reader.pages[n - 1])
    buffer = io.BytesIO()
    writer.write(buffer)
    return {"pdf_base64": base64.b64encode(buffer.getvalue()).decode("utf-8")}


def read_pdf_contents(params):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(base64.b64decode(params["pdf_base64"])))
    pages = [page.extract_text() or "" for page in reader.pages]
    return {"pages": pages, "page_count": len(pages)}


# ───────────────────────────────────────────────────────────
# Entry point — chamado via subprocess pelo Node (ver bridge.js)
# Uso: python3 documents.py <nome_da_funcao> '<json_params>'
# ───────────────────────────────────────────────────────────

if __name__ == "__main__":
    func_name = sys.argv[1]
    params = json.loads(sys.argv[2])

    funcoes = {
        "create_pdf": create_pdf,
        "create_pdf_structured": create_pdf_structured,
        "create_docx": create_docx,
        "merge_pdfs": merge_pdfs,
        "split_pdf_pages": split_pdf_pages,
        "read_pdf_contents": read_pdf_contents,
    }

    if func_name not in funcoes:
        print(json.dumps({"error": f"Função '{func_name}' não encontrada"}))
        sys.exit(1)

    try:
        resultado = funcoes[func_name](params)
        print(json.dumps(resultado))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)